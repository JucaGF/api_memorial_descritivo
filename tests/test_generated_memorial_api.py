from __future__ import annotations

import asyncio
import json
import unittest
from io import BytesIO
from datetime import datetime, timezone
from pathlib import Path
from unittest.mock import AsyncMock, MagicMock, patch

from docx import Document
from fastapi.responses import JSONResponse
from starlette.datastructures import UploadFile
from starlette.middleware.cors import CORSMiddleware

import app.api.routes as routes
from app.main import app
from app.schemas.generated_memorial import GeneratedMemorialResponse
from app.services.generated_memorial_store import (
    GeneratedMemorialArtifactNotFoundError,
    GeneratedMemorialStorageError,
)
from app.services.memorial_validator import MemorialValidationError, ValidationIssue


def _memorial(memorial_id: str = "abc-123", memorial_type: str = "telecom") -> GeneratedMemorialResponse:
    now = datetime.now(tz=timezone.utc)
    return GeneratedMemorialResponse(
        id=memorial_id,
        type=memorial_type,
        project_name="Memorial Telecom",
        status="ready",
        observations="Observacao",
        pdf_filenames=["projeto.pdf"],
        created_at=now,
        updated_at=now,
        download_url="https://signed.example/download",
    )


def _request(method: str = "GET", path: str = "/api/v1/memoriais") -> MagicMock:
    request = MagicMock()
    request.state.request_id = "req-123"
    request.method = method
    request.url.path = path
    return request


class GeneratedMemorialApiTests(unittest.TestCase):
    def test_validation_error_response_uses_quantitative_conflict_code(self) -> None:
        error = MemorialValidationError(
            issues=[
                ValidationIssue(
                    path="$.pontos_utilizacao.conflitos",
                    message="Conflitos criticos GLP v2 sem resolucao.",
                    validator="glp_v2_conflict",
                )
            ],
            extraction_report={
                "conflicts": [
                    {
                        "tipo": "glp_v2_points_total_mismatch",
                        "status": "unresolved",
                        "valores_observados": [61, 56],
                    }
                ]
            },
        )

        response = routes._validation_error_response(
            error,
            "Payload invalido para o memorial GLP v2.",
            request=_request("POST", "/api/v1/memoriais/glp/v2/from-files/persist"),
        )

        body = json.loads(response.body.decode("utf-8"))
        self.assertEqual(response.status_code, 409)
        self.assertEqual(body["detail"], body["error"]["message"])
        self.assertEqual(body["error"]["code"], "quantitative_conflict_unresolved")
        self.assertIn("valores diferentes", body["error"]["message"])
        self.assertEqual(
            body["conflicts"][0]["tipo"],
            "glp_v2_points_total_mismatch",
        )
        self.assertEqual(body["errors"][0]["validator"], "glp_v2_conflict")
        self.assertEqual(
            body["error"]["details"]["conflicts"][0]["valores_observados"],
            [61, 56],
        )
        self.assertIn("extraction_report", body)

    def test_extract_conflicts_from_report_reads_quantitative_conflicts(self) -> None:
        report = {
            "cross_validation": {
                "quantitative_conflicts": [
                    {
                        "tipo": "glp_v2_points_total_mismatch",
                        "status": "resolved",
                    }
                ]
            }
        }

        conflicts = routes._extract_conflicts_from_report(report)

        self.assertEqual(conflicts, report["cross_validation"]["quantitative_conflicts"])

    def test_cors_allows_local_vite_frontend_origin(self) -> None:
        cors_middleware = next(
            middleware
            for middleware in app.user_middleware
            if middleware.cls is CORSMiddleware
        )
        self.assertIn("http://localhost:5173", cors_middleware.kwargs["allow_origins"])
        self.assertIn("http://127.0.0.1:5173", cors_middleware.kwargs["allow_origins"])

    @patch("app.api.routes.create_generated_memorial")
    @patch(
        "app.api.routes.generate_memorial_telecom_v1_from_uploaded_files",
        new_callable=AsyncMock,
    )
    def test_post_persisted_telecom_from_files_returns_memorial_metadata(
        self,
        pipeline_mock,
        create_mock,
    ) -> None:
        async def pipeline_side_effect(_files, output_path: Path):
            document = Document()
            document.add_paragraph("Memorial telecom gerado.")
            document.save(output_path)

        pipeline_mock.side_effect = pipeline_side_effect
        create_mock.return_value = _memorial()

        files = [UploadFile(filename="projeto.pdf", file=BytesIO(b"%PDF-1.4 teste"))]
        response = asyncio.run(
            routes.create_persisted_memorial_from_files(
                "telecom",
                MagicMock(),
                files,
                "Observacao",
            )
        )

        self.assertEqual(response.id, "abc-123")
        self.assertEqual(response.type, "telecom")
        self.assertEqual(response.project_name, "Memorial Telecom")
        self.assertEqual(response.pdf_filenames, ["projeto.pdf"])
        self.assertEqual(response.download_url, "https://signed.example/download")
        create_mock.assert_called_once()
        self.assertEqual(create_mock.call_args.kwargs["memorial_type"], "telecom")
        self.assertEqual(create_mock.call_args.kwargs["observations"], "Observacao")

    @patch("app.api.routes.create_generated_memorial")
    @patch(
        "app.api.routes.generate_memorial_telecom_v1_from_uploaded_files",
        new_callable=AsyncMock,
    )
    def test_post_persisted_telecom_from_files_returns_safe_503_for_storage_failure(
        self,
        pipeline_mock,
        create_mock,
    ) -> None:
        async def pipeline_side_effect(_files, output_path: Path):
            document = Document()
            document.add_paragraph("Memorial telecom gerado.")
            document.save(output_path)

        pipeline_mock.side_effect = pipeline_side_effect
        create_mock.side_effect = GeneratedMemorialStorageError("internal storage /srv/private")

        files = [UploadFile(filename="projeto.pdf", file=BytesIO(b"%PDF-1.4 teste"))]
        request = MagicMock()
        request.state.request_id = "req-123"
        request.method = "POST"
        request.url.path = "/api/v1/memoriais/telecom/from-files/persist"

        response = asyncio.run(
            routes.create_persisted_memorial_from_files(
                "telecom",
                request,
                files,
                "Observacao",
            )
        )

        self.assertIsInstance(response, JSONResponse)
        self.assertEqual(response.status_code, 503)
        body = response.body.decode("utf-8")
        self.assertIn("Armazenamento do memorial indisponível.", body)
        self.assertNotIn("/srv/private", body)

    @patch(
        "app.api.routes.generate_memorial_gas_natural_v1_from_uploaded_files",
        new_callable=AsyncMock,
    )
    def test_post_persisted_gas_natural_from_files_logs_validation_issues_and_report(
        self,
        pipeline_mock,
    ) -> None:
        pipeline_mock.side_effect = MemorialValidationError(
            [
                ValidationIssue(
                    path="$.crm",
                    message="'pavimento' is a required property",
                    validator="required",
                ),
                ValidationIssue(
                    path="$.ramal.primario_diametro",
                    message="None is not of type 'number'",
                    validator="type",
                ),
            ],
            extraction_report={
                "filled": ["obra.construtora"],
                "missing": ["obra.nome"],
                "pending": ["crm.pavimento", "ramal.primario_diametro"],
            },
        )

        request = MagicMock()
        request.state.request_id = "req-gas-123"
        request.method = "POST"
        request.url.path = "/api/v1/memoriais/gas-natural/from-files/persist"
        files = [UploadFile(filename="projeto.pdf", file=BytesIO(b"%PDF-1.4 teste"))]

        with self.assertLogs("app.api.routes", level="WARNING") as captured:
            response = asyncio.run(
                routes.create_persisted_memorial_from_files(
                    "gas-natural",
                    request,
                    files,
                    None,
                )
            )

        self.assertIsInstance(response, JSONResponse)
        self.assertEqual(response.status_code, 400)
        log_output = "\n".join(captured.output)
        self.assertIn("$.crm", log_output)
        self.assertIn("$.ramal.primario_diametro", log_output)
        self.assertIn("\"missing_count\": 1", log_output)
        self.assertIn("\"pending_count\": 2", log_output)

    @patch("app.api.routes.list_generated_memorials")
    def test_get_memoriais_lists_persisted_memorials(self, list_mock) -> None:
        list_mock.return_value = [_memorial()]

        response = routes.list_persisted_memorials(_request(), type="telecom")

        self.assertEqual(len(response.memorials), 1)
        self.assertEqual(response.memorials[0].type, "telecom")
        list_mock.assert_called_once_with("telecom")

    @patch("app.api.routes.list_generated_memorials")
    def test_get_memoriais_accepts_glp_v2_filter(self, list_mock) -> None:
        list_mock.return_value = [_memorial(memorial_type="glp_v2")]

        response = routes.list_persisted_memorials(_request(), type="glp_v2")

        self.assertEqual(len(response.memorials), 1)
        self.assertEqual(response.memorials[0].type, "glp_v2")
        list_mock.assert_called_once_with("glp_v2")

    @patch("app.api.routes.get_generated_memorial")
    def test_get_memorial_returns_404_for_unknown_id(self, get_mock) -> None:
        get_mock.return_value = None

        response = routes.get_persisted_memorial("missing", _request())

        self.assertEqual(response.status_code, 404)
        body = json.loads(response.body.decode("utf-8"))
        self.assertEqual(body["detail"], "Memorial não encontrado.")
        self.assertEqual(body["error"]["code"], "generated_memorial_not_found")

    @patch("app.api.routes.get_generated_memorial")
    def test_get_memorial_returns_detail(self, get_mock) -> None:
        get_mock.return_value = _memorial()

        response = routes.get_persisted_memorial("abc-123", _request())

        self.assertEqual(response.id, "abc-123")

    @patch("app.api.routes.create_signed_download_url")
    @patch("app.api.routes.get_generated_memorial_record")
    def test_get_memorial_download_returns_signed_url(self, get_record_mock, signed_url_mock) -> None:
        get_record_mock.return_value = {
            "id": "abc-123",
            "type": "telecom",
            "status": "ready",
            "storage_bucket": "generated-memorials",
            "storage_path": "telecom/abc-123/memorial_telecom_v1.docx",
        }
        signed_url_mock.return_value = "https://signed.example/download"

        request = MagicMock()
        request.state.request_id = "req-123"
        request.method = "GET"
        request.url.path = "/api/v1/memoriais/abc-123/download"

        response = routes.get_persisted_memorial_download("abc-123", request)

        self.assertEqual(response.download_url, "https://signed.example/download")

    @patch("app.api.routes.create_signed_download_url")
    @patch("app.api.routes.get_generated_memorial_record")
    def test_get_memorial_download_returns_404_when_artifact_is_missing(
        self,
        get_record_mock,
        signed_url_mock,
    ) -> None:
        get_record_mock.return_value = {
            "id": "abc-123",
            "type": "telecom",
            "status": "ready",
            "storage_bucket": "generated-memorials",
            "storage_path": "telecom/abc-123/memorial_telecom_v1.docx",
        }
        signed_url_mock.side_effect = GeneratedMemorialArtifactNotFoundError("missing")

        request = MagicMock()
        request.state.request_id = "req-123"
        request.method = "GET"
        request.url.path = "/api/v1/memoriais/abc-123/download"

        response = routes.get_persisted_memorial_download("abc-123", request)

        self.assertIsInstance(response, JSONResponse)
        self.assertEqual(response.status_code, 404)
        self.assertIn("não está mais disponível", response.body.decode("utf-8"))
        self.assertNotIn("telecom/abc-123", response.body.decode("utf-8"))

    @patch("app.api.routes.create_signed_download_url")
    @patch("app.api.routes.get_generated_memorial_record")
    def test_get_memorial_download_returns_safe_503_for_storage_failure(
        self,
        get_record_mock,
        signed_url_mock,
    ) -> None:
        get_record_mock.return_value = {
            "id": "abc-123",
            "type": "telecom",
            "status": "ready",
            "storage_bucket": "generated-memorials",
            "storage_path": "telecom/abc-123/memorial_telecom_v1.docx",
        }
        signed_url_mock.side_effect = GeneratedMemorialStorageError(
            "raw internal storage path /srv/private"
        )

        request = MagicMock()
        request.state.request_id = "req-123"
        request.method = "GET"
        request.url.path = "/api/v1/memoriais/abc-123/download"

        response = routes.get_persisted_memorial_download("abc-123", request)

        self.assertIsInstance(response, JSONResponse)
        self.assertEqual(response.status_code, 503)
        body = response.body.decode("utf-8")
        self.assertIn("Armazenamento do memorial indisponível.", body)
        self.assertNotIn("/srv/private", body)

    @patch("app.api.routes.create_signed_download_url")
    @patch("app.api.routes.get_generated_memorial_record")
    def test_get_memorial_download_returns_409_when_generation_failed(
        self,
        get_record_mock,
        signed_url_mock,
    ) -> None:
        get_record_mock.return_value = {
            "id": "abc-123",
            "type": "telecom",
            "status": "failed",
            "storage_bucket": "generated-memorials",
            "storage_path": "telecom/abc-123/memorial_telecom_v1.docx",
        }

        request = MagicMock()
        request.state.request_id = "req-123"
        request.method = "GET"
        request.url.path = "/api/v1/memoriais/abc-123/download"

        response = routes.get_persisted_memorial_download("abc-123", request)

        self.assertIsInstance(response, JSONResponse)
        self.assertEqual(response.status_code, 409)
        self.assertIn("ainda não está disponível para download", response.body.decode("utf-8"))
        signed_url_mock.assert_not_called()

    @patch("app.api.routes.delete_generated_memorial")
    def test_delete_memorial_returns_204_when_deleted(self, delete_mock) -> None:
        delete_mock.return_value = True

        response = routes.delete_persisted_memorial(
            "abc-123", _request("DELETE", "/api/v1/memoriais/abc-123")
        )

        self.assertEqual(response.status_code, 204)
        delete_mock.assert_called_once_with("abc-123")

    @patch("app.api.routes.delete_generated_memorial")
    def test_delete_memorial_returns_404_when_missing(self, delete_mock) -> None:
        delete_mock.return_value = False

        response = routes.delete_persisted_memorial(
            "missing", _request("DELETE", "/api/v1/memoriais/missing")
        )

        self.assertEqual(response.status_code, 404)
        body = json.loads(response.body.decode("utf-8"))
        self.assertEqual(body["detail"], "Memorial não encontrado.")
        self.assertEqual(body["error"]["code"], "generated_memorial_not_found")

    @patch("app.api.routes.delete_generated_memorial")
    def test_delete_memorial_returns_safe_503_for_storage_failure(self, delete_mock) -> None:
        delete_mock.side_effect = GeneratedMemorialStorageError("leaked path /srv/private")

        request = MagicMock()
        request.state.request_id = "req-123"
        request.method = "DELETE"
        request.url.path = "/api/v1/memoriais/abc-123"

        response = routes.delete_persisted_memorial("abc-123", request)

        self.assertIsInstance(response, JSONResponse)
        self.assertEqual(response.status_code, 503)
        body = response.body.decode("utf-8")
        self.assertIn("Armazenamento do memorial indisponível.", body)
        self.assertNotIn("/srv/private", body)

    @patch("app.api.routes.create_generated_memorial")
    @patch(
        "app.api.routes.generate_memorial_glp_v1_from_uploaded_files",
        new_callable=AsyncMock,
    )
    def test_persist_passes_final_context_and_versions_to_store(
        self,
        pipeline_mock,
        create_mock,
    ) -> None:
        from app.services.pipeline import PipelineResult
        from app.services.extraction_mapper import ExtractionReport

        async def pipeline_side_effect(_files, output_path: Path):
            document = Document()
            document.add_paragraph("Memorial GLP gerado.")
            document.save(output_path)
            return PipelineResult(
                context={"obra": {"nome": "Exemplo"}, "abastecimento": {"qtd_tanques": 1}},
                output_path=output_path,
                extraction_report=ExtractionReport(
                    filled=["obra.nome"],
                    missing=[],
                    pending=[],
                ),
            )

        pipeline_mock.side_effect = pipeline_side_effect
        create_mock.return_value = _memorial(memorial_type="glp")

        files = [UploadFile(filename="projeto.pdf", file=BytesIO(b"%PDF-1.4 teste"))]
        asyncio.run(
            routes.create_persisted_memorial_from_files(
                "glp",
                MagicMock(),
                files,
                None,
            )
        )

        kwargs = create_mock.call_args.kwargs
        self.assertEqual(kwargs["memorial_type"], "glp")
        self.assertEqual(kwargs["context_version"], "glp_v1")
        self.assertEqual(kwargs["template_version"], "glp_v1")
        self.assertEqual(kwargs["final_context"]["abastecimento"]["qtd_tanques"], 1)
        self.assertIn("filled", kwargs["extraction_report"])
        self.assertEqual(kwargs["extraction_report"]["filled"], ["obra.nome"])
        self.assertEqual(kwargs["conflicts"], [])

    @patch("app.api.routes.get_generated_memorial")
    def test_get_memorial_with_include_context_propagates_flag(self, get_mock) -> None:
        get_mock.return_value = _memorial()

        routes.get_persisted_memorial("abc-123", _request(), include_context=True)

        get_mock.assert_called_once_with("abc-123", include_context=True)

    @patch("app.api.routes.get_generated_memorial")
    def test_get_memorial_default_does_not_request_context(self, get_mock) -> None:
        get_mock.return_value = _memorial()

        routes.get_persisted_memorial("abc-123", _request())

        get_mock.assert_called_once_with("abc-123", include_context=False)

    @patch("app.api.routes.create_generated_memorial")
    @patch("app.api.routes.generate_memorial_glp_v2")
    @patch("app.api.routes.get_generated_memorial_record")
    def test_correct_persisted_memorial_creates_new_version_for_any_type(
        self,
        get_record_mock,
        generate_mock,
        create_mock,
    ) -> None:
        from app.services.pipeline import PipelineResult

        original_created_at = datetime.now(tz=timezone.utc).isoformat()
        get_record_mock.return_value = {
            "id": "old-123",
            "type": "glp_v2",
            "project_name": "Memorial GLP v2",
            "status": "ready",
            "observations": "Observacao original",
            "pdf_filenames": ["gas.pdf"],
            "final_context": {
                "obra": {"nome": "MGAMAK", "qtd_apartamentos": 30},
                "pontos_utilizacao": {"fogao": {"quantidade": 34}},
            },
            "extraction_report": {
                "missing": [],
                "evidence": {
                    "obra.qtd_apartamentos": {
                        "value": 30,
                        "rule": "apartment_visual_labels",
                        "evidence": "APTO 001...801",
                        "confidence": "low",
                    }
                },
            },
            "conflicts": [],
            "context_version": "glp_v2",
            "template_version": "glp_v2",
            "created_at": original_created_at,
            "updated_at": original_created_at,
        }
        generate_mock.return_value = PipelineResult(
            context={
                "obra": {"nome": "MGAMAK", "qtd_apartamentos": 29},
                "pontos_utilizacao": {"fogao": {"quantidade": 35}},
            },
            output_path=Path("/tmp/memorial_glp_v2_corrigido.docx"),
        )
        create_mock.return_value = _memorial(memorial_id="new-456", memorial_type="glp_v2")

        payload = routes.MemorialCorrectionsPayload(
            corrections={
                "obra": {"qtd_apartamentos": 29},
                "pontos_utilizacao": {"fogao": {"quantidade": 35}},
            }
        )
        response = routes.correct_persisted_memorial(
            "old-123",
            payload,
            _request("POST", "/api/v1/memoriais/old-123/correcoes"),
        )

        self.assertEqual(response.id, "new-456")
        generate_mock.assert_called_once()
        corrected_context = generate_mock.call_args.args[0]
        self.assertEqual(corrected_context["obra"]["qtd_apartamentos"], 29)
        self.assertEqual(
            corrected_context["pontos_utilizacao"]["fogao"]["quantidade"],
            35,
        )
        create_kwargs = create_mock.call_args.kwargs
        self.assertEqual(create_kwargs["memorial_type"], "glp_v2")
        self.assertEqual(create_kwargs["pdf_filenames"], ["gas.pdf"])
        self.assertEqual(create_kwargs["extraction_report"]["user_corrections"]["obra.qtd_apartamentos"], 29)
        self.assertEqual(
            create_kwargs["extraction_report"]["user_corrections"]["pontos_utilizacao.fogao.quantidade"],
            35,
        )

    @patch("app.api.routes.get_generated_memorial_record")
    def test_correct_persisted_memorial_requires_stored_context(self, get_record_mock) -> None:
        get_record_mock.return_value = {
            "id": "old-123",
            "type": "telecom",
            "status": "ready",
            "final_context": None,
        }

        response = routes.correct_persisted_memorial(
            "old-123",
            routes.MemorialCorrectionsPayload(corrections={"obra": {"nome": "Novo"}}),
            _request("POST", "/api/v1/memoriais/old-123/correcoes"),
        )

        self.assertEqual(response.status_code, 409)
        body = json.loads(response.body.decode("utf-8"))
        self.assertEqual(body["error"]["code"], "generated_memorial_context_missing")


if __name__ == "__main__":
    unittest.main()
