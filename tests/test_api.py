from __future__ import annotations

import json
import importlib
from pathlib import Path
from tempfile import TemporaryDirectory
import unittest
from unittest.mock import AsyncMock, MagicMock, patch
from datetime import datetime, timedelta, timezone

from docx import Document
from fastapi.testclient import TestClient

from app.main import app
from app.services.extraction_mapper import ExtractionReport, MappingResult
from app.services.memorial_validator import MemorialValidationError, ValidationIssue
from app.services.pipeline import PipelineResult
from app.services.session_store import ReviewSession


ROOT = Path(__file__).resolve().parent.parent
FIXTURES_DIR = ROOT / "tests" / "fixtures"
PROJECTS_DIR = ROOT / "projects"


def load_fixture(filename: str) -> dict:
    with (FIXTURES_DIR / filename).open("r", encoding="utf-8") as file:
        return json.load(file)


class ApiTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        cls.client = TestClient(app)

    def test_health_live_returns_200_with_safe_json(self) -> None:
        response = self.client.get("/health/live")

        self.assertEqual(response.status_code, 200)
        body = response.json()
        self.assertEqual(body["status"], "ok")
        self.assertIn("timestamp", body)
        self.assertEqual(len(body["checks"]), 1)
        self.assertEqual(body["checks"][0]["name"], "app")
        self.assertEqual(body["checks"][0]["status"], "ok")

    @patch(
        "app.api.routes.get_readiness_payload",
        return_value={
            "status": "ok",
            "timestamp": "2026-04-30T12:00:00+00:00",
            "checks": [
                {"name": "app", "status": "ok"},
                {"name": "templates", "status": "ok"},
                {"name": "storage", "status": "ok"},
                {"name": "configuration", "status": "ok"},
            ],
        },
        create=True,
    )
    def test_health_ready_returns_200_when_checks_pass(self, _) -> None:
        response = self.client.get("/health/ready")

        self.assertEqual(response.status_code, 200)
        body = response.json()
        self.assertEqual(body["status"], "ok")
        self.assertEqual(
            [check["name"] for check in body["checks"]],
            ["app", "templates", "storage", "configuration"],
        )

    @patch(
        "app.api.routes.get_readiness_payload",
        return_value={
            "status": "error",
            "timestamp": "2026-04-30T12:00:00+00:00",
            "checks": [
                {"name": "app", "status": "ok"},
                {"name": "templates", "status": "error"},
            ],
        },
        create=True,
    )
    def test_health_ready_returns_503_when_critical_check_fails(self, _) -> None:
        response = self.client.get("/health/ready")

        self.assertEqual(response.status_code, 503)
        body = response.json()
        self.assertEqual(body["status"], "error")
        self.assertEqual(body["checks"][1]["name"], "templates")
        self.assertEqual(body["checks"][1]["status"], "error")

    @patch(
        "app.api.routes.get_readiness_payload",
        return_value={
            "status": "ok",
            "timestamp": "2026-04-30T12:00:00+00:00",
            "checks": [
                {
                    "name": "configuration",
                    "status": "ok",
                    "detail": "configured",
                }
            ],
        },
        create=True,
    )
    def test_health_endpoints_do_not_expose_sensitive_values(self, _) -> None:
        with patch.dict(
            "os.environ",
            {
                "OPENAI_API_KEY": "super-secret-openai",
                "SUPABASE_SERVICE_ROLE_KEY": "super-secret-supabase",
                "DATABASE_URL": "postgres://user:password@example.com/db",
            },
            clear=False,
        ):
            live_response = self.client.get("/health/live")
            ready_response = self.client.get("/health/ready")

        self.assertEqual(live_response.status_code, 200)
        self.assertEqual(ready_response.status_code, 200)

        combined_body = json.dumps(
            {
                "live": live_response.json(),
                "ready": ready_response.json(),
            },
            ensure_ascii=False,
        )
        self.assertNotIn("super-secret-openai", combined_body)
        self.assertNotIn("super-secret-supabase", combined_body)
        self.assertNotIn("postgres://user:password@example.com/db", combined_body)

    def test_post_memorial_eletrico_returns_docx_for_valid_payload(self) -> None:
        payload = load_fixture("eletrico_com_subestacao.json")

        response = self.client.post("/api/v1/memoriais/eletrico", json=payload)

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.headers["content-type"],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        self.assertTrue(
            response.headers["content-disposition"].startswith("attachment;")
        )
        self.assertTrue(response.content.startswith(b"PK"))

    def test_post_memorial_eletrico_returns_400_for_invalid_payload(self) -> None:
        payload = load_fixture("eletrico_sem_subestacao.json")
        del payload["obra"]["nome"]

        response = self.client.post("/api/v1/memoriais/eletrico", json=payload)

        self.assertEqual(response.status_code, 400)
        body = response.json()
        self.assertEqual(
            body["detail"],
            "Payload invalido para o memorial eletrico v1.",
        )
        self.assertTrue(body["errors"])
        self.assertEqual(body["errors"][0]["path"], "$.obra")
        self.assertIn("nome", body["errors"][0]["message"])

    def test_post_memorial_telecom_returns_docx_for_valid_payload(self) -> None:
        payload = load_fixture("telecom_base.json")

        response = self.client.post("/api/v1/memoriais/telecom", json=payload)

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.headers["content-type"],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        self.assertTrue(
            response.headers["content-disposition"].startswith("attachment;")
        )
        self.assertTrue(response.content.startswith(b"PK"))

    def test_post_memorial_telecom_returns_400_for_invalid_payload(self) -> None:
        payload = load_fixture("telecom_base.json")
        del payload["obra"]["nome"]

        response = self.client.post("/api/v1/memoriais/telecom", json=payload)

        self.assertEqual(response.status_code, 400)
        body = response.json()
        self.assertEqual(
            body["detail"],
            "Payload invalido para o memorial telecom v1.",
        )
        self.assertTrue(body["errors"])
        self.assertEqual(body["errors"][0]["path"], "$.obra")
        self.assertIn("nome", body["errors"][0]["message"])

    def test_post_memorial_gas_natural_returns_docx_for_valid_payload(self) -> None:
        payload = load_fixture("gas_natural_base.json")

        response = self.client.post("/api/v1/memoriais/gas-natural", json=payload)

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.headers["content-type"],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        self.assertTrue(
            response.headers["content-disposition"].startswith("attachment;")
        )
        self.assertTrue(response.content.startswith(b"PK"))

    def test_post_memorial_gas_natural_returns_400_for_invalid_payload(self) -> None:
        payload = load_fixture("gas_natural_base.json")
        del payload["valvula"]["esfera_diametro"]

        response = self.client.post("/api/v1/memoriais/gas-natural", json=payload)

        self.assertEqual(response.status_code, 400)
        body = response.json()
        self.assertEqual(
            body["detail"],
            "Payload invalido para o memorial gas natural v1.",
        )
        self.assertTrue(body["errors"])
        self.assertEqual(body["errors"][0]["path"], "$.valvula")
        self.assertIn("esfera_diametro", body["errors"][0]["message"])

    def test_post_memorial_eletrico_upload_returns_metadata_for_valid_files(self) -> None:
        files = [
            ("files", ("projeto.pdf", b"%PDF-1.4 teste", "application/pdf")),
            (
                "files",
                (
                    "memorial.docx",
                    b"PK\x03\x04conteudo",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            ),
        ]

        response = self.client.post("/api/v1/memoriais/eletrico/upload", files=files)

        self.assertEqual(response.status_code, 200)
        body = response.json()
        self.assertEqual(len(body["files"]), 2)
        self.assertNotIn("request_dir", body)
        self.assertNotIn("saved_path", body["files"][0])
        self.assertEqual(body["files"][0]["extension"], ".pdf")
        self.assertEqual(body["files"][1]["extension"], ".docx")
        self.assertEqual(body["files"][0]["filename"], "projeto.pdf")

    def test_post_memorial_telecom_upload_returns_metadata_for_valid_files(self) -> None:
        files = [
            ("files", ("projeto.pdf", b"%PDF-1.4 teste", "application/pdf")),
            (
                "files",
                (
                    "memorial.docx",
                    b"PK\x03\x04conteudo",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            ),
        ]

        response = self.client.post("/api/v1/memoriais/telecom/upload", files=files)

        self.assertEqual(response.status_code, 200)
        body = response.json()
        self.assertEqual(len(body["files"]), 2)
        self.assertEqual(body["files"][0]["extension"], ".pdf")
        self.assertEqual(body["files"][1]["extension"], ".docx")
        self.assertEqual(body["files"][0]["filename"], "projeto.pdf")

    def test_post_memorial_gas_natural_upload_returns_metadata_for_valid_files(self) -> None:
        files = [
            ("files", ("projeto.pdf", b"%PDF-1.4 teste", "application/pdf")),
            (
                "files",
                (
                    "memorial.docx",
                    b"PK\x03\x04conteudo",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            ),
        ]

        response = self.client.post("/api/v1/memoriais/gas-natural/upload", files=files)

        self.assertEqual(response.status_code, 200)
        body = response.json()
        self.assertEqual(len(body["files"]), 2)
        self.assertEqual(body["files"][0]["extension"], ".pdf")
        self.assertEqual(body["files"][1]["extension"], ".docx")
        self.assertEqual(body["files"][0]["filename"], "projeto.pdf")

    def test_post_memorial_eletrico_upload_returns_400_for_invalid_extension(self) -> None:
        files = [
            (
                "files",
                (
                    "projeto.txt",
                    b"conteudo",
                    "text/plain",
                ),
            )
        ]

        response = self.client.post("/api/v1/memoriais/eletrico/upload", files=files)

        self.assertEqual(response.status_code, 400)
        body = response.json()
        self.assertIn("Extensao nao suportada", body["detail"])

    def test_post_memorial_eletrico_upload_returns_400_for_empty_list(self) -> None:
        response = self.client.post("/api/v1/memoriais/eletrico/upload", files=[])

        self.assertEqual(response.status_code, 400)
        self.assertEqual(
            response.json()["detail"],
            "Envie ao menos um arquivo PDF ou DOCX.",
        )

    @patch(
        "app.api.routes.generate_memorial_eletrico_v1_from_uploaded_files",
        new_callable=AsyncMock,
    )
    def test_post_memorial_eletrico_from_files_returns_docx_for_valid_files(
        self,
        pipeline_mock,
    ) -> None:
        async def pipeline_side_effect(_files, output_path: Path) -> PipelineResult:
            document = Document()
            document.add_paragraph("Memorial gerado a partir de arquivos.")
            document.save(output_path)
            return PipelineResult(
                context={"obra": {"nome": "Edificio Exemplo"}},
                output_path=output_path,
            )

        pipeline_mock.side_effect = pipeline_side_effect

        with TemporaryDirectory() as temp_dir:
            files = [
                (
                    "files",
                    (
                        "projeto.docx",
                        b"PK\x03\x04conteudo",
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                )
            ]

            response = self.client.post(
                "/api/v1/memoriais/eletrico/from-files",
                files=files,
            )

            self.assertEqual(response.status_code, 200)
            self.assertEqual(
                response.headers["content-type"],
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            self.assertTrue(response.content.startswith(b"PK"))

        pipeline_mock.assert_awaited_once()

    @patch(
        "app.api.routes.generate_memorial_telecom_v1_from_uploaded_files",
        new_callable=AsyncMock,
    )
    def test_post_memorial_telecom_from_files_returns_docx_for_valid_files(
        self,
        pipeline_mock,
    ) -> None:
        async def pipeline_side_effect(_files, output_path: Path) -> PipelineResult:
            document = Document()
            document.add_paragraph("Memorial telecom gerado a partir de arquivos.")
            document.save(output_path)
            return PipelineResult(
                context={"obra": {"nome": "Edificio Exemplo"}},
                output_path=output_path,
            )

        pipeline_mock.side_effect = pipeline_side_effect

        files = [
            (
                "files",
                (
                    "projeto.docx",
                    b"PK\x03\x04conteudo",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            )
        ]

        response = self.client.post(
            "/api/v1/memoriais/telecom/from-files",
            files=files,
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.headers["content-type"],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        self.assertTrue(response.content.startswith(b"PK"))

        pipeline_mock.assert_awaited_once()

    @patch(
        "app.api.routes.generate_memorial_gas_natural_v1_from_uploaded_files",
        new_callable=AsyncMock,
    )
    def test_post_memorial_gas_natural_from_files_returns_docx_for_valid_files(
        self,
        pipeline_mock,
    ) -> None:
        async def pipeline_side_effect(_files, output_path: Path) -> PipelineResult:
            document = Document()
            document.add_paragraph("Memorial gas natural gerado a partir de arquivos.")
            document.save(output_path)
            return PipelineResult(
                context={"obra": {"nome": "Edificio Exemplo"}},
                output_path=output_path,
            )

        pipeline_mock.side_effect = pipeline_side_effect

        files = [
            (
                "files",
                (
                    "projeto.docx",
                    b"PK\x03\x04conteudo",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            )
        ]

        response = self.client.post(
            "/api/v1/memoriais/gas-natural/from-files",
            files=files,
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.headers["content-type"],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        self.assertTrue(response.content.startswith(b"PK"))

        pipeline_mock.assert_awaited_once()

    def test_post_memorial_eletrico_from_files_returns_400_for_invalid_extension(self) -> None:
        files = [
            (
                "files",
                (
                    "projeto.txt",
                    b"conteudo",
                    "text/plain",
                ),
            )
        ]

        response = self.client.post(
            "/api/v1/memoriais/eletrico/from-files",
            files=files,
        )

        self.assertEqual(response.status_code, 400)
        self.assertIn("Extensao nao suportada", response.json()["detail"])

    @patch(
        "app.api.routes.generate_memorial_eletrico_v1_from_uploaded_files",
        new_callable=AsyncMock,
    )
    def test_post_memorial_eletrico_from_files_returns_400_for_validation_error(
        self,
        pipeline_mock,
    ) -> None:
        pipeline_mock.side_effect = MemorialValidationError(
            [
                ValidationIssue(
                    path="$",
                    message="'documento' is a required property",
                    validator="required",
                )
            ]
        )
        files = [
            (
                "files",
                (
                    "projeto.docx",
                    b"PK\x03\x04conteudo",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            )
        ]

        response = self.client.post(
            "/api/v1/memoriais/eletrico/from-files",
            files=files,
        )

        self.assertEqual(response.status_code, 400)
        body = response.json()
        self.assertEqual(body["detail"], "Payload invalido para o memorial eletrico v1.")
        self.assertTrue(body["errors"])
        self.assertEqual(body["errors"][0]["path"], "$")

    @patch(
        "app.api.routes.generate_memorial_telecom_v1_from_uploaded_files",
        new_callable=AsyncMock,
    )
    def test_post_memorial_telecom_from_files_returns_400_for_validation_error(
        self,
        pipeline_mock,
    ) -> None:
        pipeline_mock.side_effect = MemorialValidationError(
            [
                ValidationIssue(
                    path="$.obra",
                    message="'tipologia' is a required property",
                    validator="required",
                )
            ]
        )
        files = [
            (
                "files",
                (
                    "projeto.docx",
                    b"PK\x03\x04conteudo",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            )
        ]

        response = self.client.post(
            "/api/v1/memoriais/telecom/from-files",
            files=files,
        )

        self.assertEqual(response.status_code, 400)
        body = response.json()
        self.assertEqual(body["detail"], "Payload invalido para o memorial telecom v1.")
        self.assertTrue(body["errors"])
        self.assertEqual(body["errors"][0]["path"], "$.obra")

    @patch(
        "app.api.routes.generate_memorial_gas_natural_v1_from_uploaded_files",
        new_callable=AsyncMock,
    )
    def test_post_memorial_gas_natural_from_files_returns_400_for_validation_error(
        self,
        pipeline_mock,
    ) -> None:
        pipeline_mock.side_effect = MemorialValidationError(
            [
                ValidationIssue(
                    path="$.crm",
                    message="'pavimento' is a required property",
                    validator="required",
                )
            ]
        )
        files = [
            (
                "files",
                (
                    "projeto.docx",
                    b"PK\x03\x04conteudo",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            )
        ]

        response = self.client.post(
            "/api/v1/memoriais/gas-natural/from-files",
            files=files,
        )

        self.assertEqual(response.status_code, 400)
        body = response.json()
        self.assertEqual(body["detail"], "Payload invalido para o memorial gas natural v1.")
        self.assertTrue(body["errors"])
        self.assertEqual(body["errors"][0]["path"], "$.crm")

    def test_post_memorial_glp_from_real_project_files_requires_llm_extraction(self) -> None:
        project_dir = PROJECTS_DIR / "gas-glp"
        pdf_paths = sorted(project_dir.glob("*.pdf"))
        self.assertTrue(pdf_paths, "Expected PDF fixtures in projects/gas-glp")

        file_handles = []
        try:
            for pdf_path in pdf_paths:
                file_handles.append(("files", (pdf_path.name, pdf_path.open("rb"), "application/pdf")))

            with patch.dict("os.environ", {"USE_LLM_EXTRACTION": ""}):
                response = self.client.post(
                    "/api/v1/memoriais/glp/from-files",
                    files=file_handles,
                )
        finally:
            for _, (_, file_obj, _) in file_handles:
                file_obj.close()

        self.assertEqual(response.status_code, 400, response.text)
        self.assertIn("LLM", response.json()["detail"])

    @patch("app.services.pipeline_from_files.extract_glp_with_llm")
    def test_post_memorial_glp_from_real_project_files_returns_docx_with_llm_context(
        self,
        llm_mock,
    ) -> None:
        project_dir = PROJECTS_DIR / "gas-glp"
        pdf_paths = sorted(project_dir.glob("*.pdf"))
        self.assertTrue(pdf_paths, "Expected PDF fixtures in projects/gas-glp")

        payload = load_fixture("glp_base.json")
        llm_mock.return_value = {
            "obra": {
                "construtora": payload["obra"]["construtora"],
                "nome": payload["obra"]["nome"],
                "localizacao": payload["obra"]["localizacao"],
                "numero_cadastro": payload["obra"]["numero_cadastro"],
                "qtd_apartamentos": payload["obra"]["qtd_apartamentos"],
                "tipo_edificacao": None,
                "tipologia": None,
                "qtd_lojas": None,
                "qtd_restaurantes": None,
            },
            "abastecimento": payload["abastecimento"],
            "dimensionamento": {
                "qtd_fogao": payload["dimensionamento"]["qtd_fogao"],
                "qtd_aquecedor": None,
                "qtd_churrasqueira": payload["dimensionamento"]["qtd_churrasqueira"],
            },
            "soma": payload["soma"],
            "ramal": payload["ramal"],
            "numero": payload["numero"],
            "teto_ou_piso": payload["teto_ou_piso"],
        }

        file_handles = []
        try:
            for pdf_path in pdf_paths:
                file_handles.append(("files", (pdf_path.name, pdf_path.open("rb"), "application/pdf")))

            with patch.dict("os.environ", {"USE_LLM_EXTRACTION": "true"}):
                response = self.client.post(
                    "/api/v1/memoriais/glp/from-files",
                    files=file_handles,
                )
        finally:
            for _, (_, file_obj, _) in file_handles:
                file_obj.close()

        self.assertEqual(response.status_code, 200, response.text)
        self.assertEqual(
            response.headers["content-type"],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        self.assertTrue(response.content.startswith(b"PK"))


def _build_pending_session(session_id: str) -> ReviewSession:
    from datetime import datetime, timedelta, timezone
    now = datetime.now(tz=timezone.utc)
    return ReviewSession(
        session_id=session_id,
        status="pending_review",
        created_at=now.isoformat(),
        expires_at=(now + timedelta(hours=24)).isoformat(),
        partial_context={"obra": {"nome": "Makai", "construtora": "MGA LTDA"}},
        extraction_report={"filled": ["obra.nome"], "missing": ["obra.tipo_edificacao"], "pending": [], "evidence": {}},
        corrections={},
    )


class ReviewSessionApiTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        cls.client = TestClient(app)

    @patch("app.api.routes._process_review_session")
    @patch("app.api.routes.create_session", return_value="test-session-id")
    @patch("app.api.routes.ingest_uploaded_files", new_callable=AsyncMock)
    def test_post_sessoes_returns_202_with_session_id(
        self, ingest_mock, create_mock, process_mock
    ) -> None:
        ingest_mock.return_value = MagicMock(files=[])

        files = [("files", ("projeto.pdf", b"%PDF-1.4 teste", "application/pdf"))]
        response = self.client.post("/api/v1/memoriais/eletrico/sessoes", files=files)

        self.assertEqual(response.status_code, 202)
        body = response.json()
        self.assertEqual(body["session_id"], "test-session-id")
        self.assertEqual(body["status"], "processing")

    @patch("app.api.routes.cleanup_ingestion_result")
    @patch("app.api.routes.create_session", side_effect=RuntimeError("db down"))
    @patch("app.api.routes.ingest_uploaded_files", new_callable=AsyncMock)
    def test_post_sessoes_cleans_temp_files_when_create_session_fails(
        self,
        ingest_mock,
        create_mock,
        cleanup_mock,
    ) -> None:
        ingestion_result = MagicMock(files=[])
        ingest_mock.return_value = ingestion_result
        files = [("files", ("projeto.pdf", b"%PDF-1.4 teste", "application/pdf"))]

        with self.assertRaises(RuntimeError):
            self.client.post("/api/v1/memoriais/eletrico/sessoes", files=files)

        cleanup_mock.assert_called_once_with(ingestion_result)

    @patch("app.api.routes.delete_session")
    @patch("app.api.routes.cleanup_ingestion_result")
    @patch("app.api.routes.create_session", return_value="test-session-id")
    @patch("app.api.routes.ingest_uploaded_files", new_callable=AsyncMock)
    def test_post_sessoes_cleans_temp_files_and_deletes_session_when_add_task_fails(
        self,
        ingest_mock,
        create_mock,
        cleanup_mock,
        delete_mock,
    ) -> None:
        ingestion_result = MagicMock(files=[])
        ingest_mock.return_value = ingestion_result
        files = [("files", ("projeto.pdf", b"%PDF-1.4 teste", "application/pdf"))]

        with patch("app.api.routes.BackgroundTasks.add_task", side_effect=RuntimeError("task queue down")):
            with self.assertRaises(RuntimeError):
                self.client.post("/api/v1/memoriais/eletrico/sessoes", files=files)

        cleanup_mock.assert_called_once_with(ingestion_result)
        delete_mock.assert_called_once_with("test-session-id")

    @patch("app.api.routes.load_session", return_value=None)
    def test_get_sessoes_returns_404_for_unknown_id(self, _) -> None:
        response = self.client.get("/api/v1/memoriais/eletrico/sessoes/nao-existe")

        self.assertEqual(response.status_code, 404)
        self.assertIn("não encontrada", response.json()["detail"])

    def test_get_sessoes_returns_404_for_expired_session_in_filesystem_store(self) -> None:
        from app.services import session_store

        with TemporaryDirectory() as temp_dir:
            session_id = "expired-session"
            expired_session = ReviewSession(
                session_id=session_id,
                status="pending_review",
                created_at=(datetime.now(tz=timezone.utc) - timedelta(days=2)).isoformat(),
                expires_at=(datetime.now(tz=timezone.utc) - timedelta(minutes=1)).isoformat(),
            )
            session_file = Path(temp_dir) / f"{session_id}.json"
            with patch.dict("os.environ", {"SUPABASE_URL": "", "SUPABASE_KEY": ""}):
                filesystem_store = importlib.reload(session_store)
                with patch("app.services.session_store._sessions_dir", return_value=Path(temp_dir)):
                    filesystem_store.save_session(expired_session)

                    with patch("app.api.routes.load_session", filesystem_store.load_session):
                        response = self.client.get(
                            f"/api/v1/memoriais/eletrico/sessoes/{session_id}"
                        )

            self.assertEqual(response.status_code, 404)
            self.assertFalse(session_file.exists())

    @patch("app.api.routes.load_session")
    def test_get_sessoes_returns_full_session_state(self, load_mock) -> None:
        session = _build_pending_session("abc-123")
        load_mock.return_value = session

        response = self.client.get("/api/v1/memoriais/eletrico/sessoes/abc-123")

        self.assertEqual(response.status_code, 200)
        body = response.json()
        self.assertEqual(body["session_id"], "abc-123")
        self.assertEqual(body["status"], "pending_review")
        self.assertIn("partial_context", body)
        self.assertIn("extraction_report", body)
        self.assertIn("filled", body["extraction_report"])
        self.assertIn("missing", body["extraction_report"])
        self.assertIn("pending", body["extraction_report"])
        self.assertIn("evidence", body["extraction_report"])

    @patch("app.api.routes.load_session")
    def test_get_sessoes_preserves_typed_extraction_report_shape_with_evidence(self, load_mock) -> None:
        session = _build_pending_session("abc-123")
        session.extraction_report = {
            "filled": ["obra.nome"],
            "missing": ["obra.tipo_edificacao"],
            "pending": [],
            "evidence": {
                "obra.nome": {
                    "value": "Makai",
                    "rule": "carimbo_line_before_company",
                    "evidence": "'MAKAI' — linha anterior a 'MGA CONSTRUÇÕES'",
                    "confidence": "high",
                }
            },
        }
        load_mock.return_value = session

        response = self.client.get("/api/v1/memoriais/eletrico/sessoes/abc-123")

        self.assertEqual(response.status_code, 200)
        body = response.json()
        self.assertEqual(body["extraction_report"]["filled"], ["obra.nome"])
        self.assertEqual(
            body["extraction_report"]["evidence"]["obra.nome"]["rule"],
            "carimbo_line_before_company",
        )
        self.assertEqual(
            body["extraction_report"]["evidence"]["obra.nome"]["confidence"],
            "high",
        )

    @patch("app.api.routes.load_session")
    def test_get_sessoes_accepts_empty_extraction_report_for_processing_compatibility(
        self,
        load_mock,
    ) -> None:
        now = datetime.now(tz=timezone.utc)
        load_mock.return_value = ReviewSession(
            session_id="abc-123",
            status="processing",
            created_at=now.isoformat(),
            expires_at=(now + timedelta(hours=24)).isoformat(),
            extraction_report={},
        )

        response = self.client.get("/api/v1/memoriais/eletrico/sessoes/abc-123")

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["extraction_report"], {})

    @patch("app.api.routes.load_session")
    def test_patch_contexto_returns_409_when_still_processing(self, load_mock) -> None:
        from datetime import datetime, timedelta, timezone
        now = datetime.now(tz=timezone.utc)
        load_mock.return_value = ReviewSession(
            session_id="abc-123",
            status="processing",
            created_at=now.isoformat(),
            expires_at=(now + timedelta(hours=24)).isoformat(),
        )

        response = self.client.patch(
            "/api/v1/memoriais/eletrico/sessoes/abc-123/contexto",
            json={"corrections": {"obra": {"nome": "Novo Nome"}}},
        )

        self.assertEqual(response.status_code, 409)

    @patch("app.api.routes.update_session")
    @patch("app.api.routes.load_session")
    def test_patch_contexto_accumulates_corrections(
        self, load_mock, update_mock
    ) -> None:
        session = _build_pending_session("abc-123")
        session_with_existing = ReviewSession(
            **{**session.__dict__, "corrections": {"obra": {"tipo_edificacao": "Residencial"}}}
        )
        load_mock.return_value = session_with_existing

        updated_session = ReviewSession(
            **{**session_with_existing.__dict__,
               "corrections": {"obra": {"tipo_edificacao": "Residencial", "tipologia": "Vertical"}}}
        )
        update_mock.return_value = updated_session

        response = self.client.patch(
            "/api/v1/memoriais/eletrico/sessoes/abc-123/contexto",
            json={"corrections": {"obra": {"tipologia": "Vertical"}}},
        )

        self.assertEqual(response.status_code, 200)
        update_mock.assert_called_once()
        merged = update_mock.call_args[1]["corrections"]
        self.assertEqual(merged["obra"]["tipo_edificacao"], "Residencial")
        self.assertEqual(merged["obra"]["tipologia"], "Vertical")

    @patch("app.api.routes.load_session", return_value=None)
    def test_post_gerar_returns_404_for_unknown_session(self, _) -> None:
        response = self.client.post("/api/v1/memoriais/eletrico/sessoes/nao-existe/gerar")

        self.assertEqual(response.status_code, 404)

    @patch("app.api.routes.load_session")
    def test_post_gerar_returns_409_when_session_still_processing(self, load_mock) -> None:
        from datetime import datetime, timedelta, timezone
        now = datetime.now(tz=timezone.utc)
        load_mock.return_value = ReviewSession(
            session_id="abc-123",
            status="processing",
            created_at=now.isoformat(),
            expires_at=(now + timedelta(hours=24)).isoformat(),
        )

        response = self.client.post("/api/v1/memoriais/eletrico/sessoes/abc-123/gerar")

        self.assertEqual(response.status_code, 409)

    @patch("app.api.routes.delete_session")
    @patch("app.api.routes.generate_memorial_eletrico_v1")
    @patch("app.api.routes.load_session")
    def test_post_gerar_returns_docx_on_success(
        self, load_mock, generate_mock, delete_mock
    ) -> None:
        session = _build_pending_session("abc-123")
        load_mock.return_value = session

        def generate_side_effect(context, output_path):
            doc = Document()
            doc.add_paragraph("Memorial gerado.")
            doc.save(output_path)
            return PipelineResult(context=context, output_path=output_path)

        generate_mock.side_effect = generate_side_effect

        response = self.client.post("/api/v1/memoriais/eletrico/sessoes/abc-123/gerar")

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.headers["content-type"],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        self.assertTrue(response.content.startswith(b"PK"))

    @patch("app.api.routes.generate_memorial_eletrico_v1")
    @patch("app.api.routes.load_session")
    def test_post_gerar_returns_400_on_validation_error(
        self, load_mock, generate_mock
    ) -> None:
        session = _build_pending_session("abc-123")
        load_mock.return_value = session
        generate_mock.side_effect = MemorialValidationError(
            [ValidationIssue(path="$.obra", message="'tipo_edificacao' is a required property", validator="required")]
        )

        response = self.client.post("/api/v1/memoriais/eletrico/sessoes/abc-123/gerar")

        self.assertEqual(response.status_code, 400)
        body = response.json()
        self.assertEqual(body["detail"], "Payload invalido para o memorial eletrico v1.")
        self.assertTrue(body["errors"])


if __name__ == "__main__":
    unittest.main()
