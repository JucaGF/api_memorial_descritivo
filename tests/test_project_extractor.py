from __future__ import annotations

from pathlib import Path
from tempfile import TemporaryDirectory
import unittest

from docx import Document

from app.services.file_ingestion import IngestedFileMetadata
from app.services.project_extractor import (
    ProjectExtractionError,
    extract_project_files,
    has_pdf_extractor_dependency,
)


class ProjectExtractorTests(unittest.TestCase):
    def test_extract_project_files_from_docx(self) -> None:
        with TemporaryDirectory() as temp_dir:
            file_path = Path(temp_dir) / "projeto.docx"
            document = Document()
            document.add_paragraph("Quadro geral de baixa tensão.")
            document.add_paragraph("Gerador de emergência parcial.")
            document.save(file_path)

            files = [
                IngestedFileMetadata(
                    original_filename="projeto.docx",
                    stored_filename="01_projeto.docx",
                    content_type=(
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    ),
                    extension=".docx",
                    size_bytes=file_path.stat().st_size,
                    saved_path=str(file_path),
                )
            ]

            result = extract_project_files(files)

            self.assertIn("Quadro geral de baixa tensão.", result.raw_text)
            self.assertEqual(result.signals["total_files"], 1)
            self.assertTrue(result.signals["has_docx"])
            self.assertFalse(result.signals["has_pdf"])

    @unittest.skipUnless(
        has_pdf_extractor_dependency(),
        "PyMuPDF nao esta instalado na .venv atual",
    )
    def test_extract_project_files_from_pdf(self) -> None:
        import fitz

        with TemporaryDirectory() as temp_dir:
            file_path = Path(temp_dir) / "projeto.pdf"
            document = fitz.open()
            page = document.new_page()
            page.insert_text((72, 72), "Entrada de energia em media tensao")
            document.save(file_path)
            document.close()

            files = [
                IngestedFileMetadata(
                    original_filename="projeto.pdf",
                    stored_filename="01_projeto.pdf",
                    content_type="application/pdf",
                    extension=".pdf",
                    size_bytes=file_path.stat().st_size,
                    saved_path=str(file_path),
                )
            ]

            result = extract_project_files(files)

            self.assertIn("Entrada de energia", result.raw_text)
            self.assertTrue(result.signals["has_pdf"])

    def test_extract_project_files_raises_for_corrupted_docx(self) -> None:
        with TemporaryDirectory() as temp_dir:
            file_path = Path(temp_dir) / "corrompido.docx"
            file_path.write_bytes(b"conteudo invalido para docx")

            files = [
                IngestedFileMetadata(
                    original_filename="corrompido.docx",
                    stored_filename="01_corrompido.docx",
                    content_type=(
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    ),
                    extension=".docx",
                    size_bytes=file_path.stat().st_size,
                    saved_path=str(file_path),
                )
            ]

            with self.assertRaises(ProjectExtractionError) as error_info:
                extract_project_files(files)

            self.assertIn("arquivo corrompido ou ilegivel", str(error_info.exception))


if __name__ == "__main__":
    unittest.main()
